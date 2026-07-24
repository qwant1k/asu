# -*- coding: utf-8 -*-
"""Генерация документа «Описание базы данных» ИС «АСУ»
по ГОСТ 34.321-96 и комплексу стандартов ГОСТ 34 (применяются в Республике Казахстан).
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from db_doc_data import ALL_SECTIONS
from generate_db_doc import (
    diagram_overview, diagram_org, diagram_refs, diagram_stock,
    diagram_requests, diagram_documents, diagram_notifications, diagram_statuses,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

FONT = 'Times New Roman'


# ============================================================
# Помощники оформления по ГОСТ
# ============================================================

def _set_run(run, size=14, bold=False, italic=False, caps=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if caps:
        run.font.all_caps = True
    run.font.color.rgb = RGBColor(0, 0, 0)


def para(doc, text='', size=14, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, caps=False,
         space_after=0, line=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if indent:
        pf.first_line_indent = Cm(1.25)
    if text:
        run = p.add_run(text)
        _set_run(run, size, bold, italic, caps)
    return p


def heading(doc, text, level=1, numbered=True):
    """Заголовок раздела по ГОСТ: полужирный, с абзацного отступа.
    Использует встроенные стили Heading для работы автособираемого содержания."""
    p = doc.add_paragraph(style=f'Heading {min(level, 3)}')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.25)
    pf.keep_with_next = True
    run = p.add_run(text)
    _set_run(run, 14 if level > 1 else 14, bold=True, caps=(level == 1 and not numbered))
    return p


def structural_heading(doc, text):
    """Структурный заголовок (АННОТАЦИЯ, СОДЕРЖАНИЕ...) — по центру, прописными."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(12)
    run = p.add_run(text)
    _set_run(run, 14, bold=True, caps=True)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    _set_run(run, size, bold)


class Counters:
    def __init__(self):
        self.table = 0
        self.figure = 0


def table_caption(doc, counters, title):
    counters.table += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f'Таблица {counters.table} — {title}')
    _set_run(run, 13)
    return counters.table


def figure_caption(doc, counters, title):
    counters.figure += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f'Рисунок {counters.figure} — {title}')
    _set_run(run, 13)
    return counters.figure


def add_image(doc, path, width_cm=16.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Cm(width_cm))


def fields_table(doc, fields):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, name in enumerate(['Наименование поля', 'Тип данных / связь', 'Назначение']):
        cell_text(hdr[i], name, 12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[i], 'D9E2F3')
    for fname, ftype, fdesc in fields:
        row = table.add_row().cells
        cell_text(row[0], fname, 12)
        cell_text(row[1], ftype, 12)
        cell_text(row[2], fdesc, 12)
    widths = [Cm(4.6), Cm(5.4), Cm(6.8)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
    return table


def add_page_number_field(doc):
    """Номер страницы вверху по центру (ГОСТ 2.105 / 7.32)."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run = p.add_run()
    _set_run(run, 12)
    run._element.append(fld_begin)
    run2 = p.add_run()
    _set_run(run2, 12)
    run2._element.append(instr)
    run3 = p.add_run()
    _set_run(run3, 12)
    run3._element.append(fld_end)
    section.different_first_page_header_footer = True


def add_toc_field(doc):
    """Поле автоматического содержания (обновляется в Word: Ctrl+A, F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    _set_run(run, 14)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t')
    t.text = 'Содержание обновляется автоматически: выделите его и нажмите F9.'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r = run._element
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(t)
    r.append(fld_end)


# ============================================================
# Построение документа
# ============================================================

def build(images):
    doc = Document()
    counters = Counters()

    # Базовый стиль
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(0)

    # Стили заголовков по ГОСТ: Times New Roman, чёрный, полужирный
    for lvl in (1, 2, 3):
        h = doc.styles[f'Heading {lvl}']
        h.font.name = FONT
        h.font.size = Pt(14)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # Поля страницы по ГОСТ: левое 30, правое 10, верхнее и нижнее 20 мм
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    add_page_number_field(doc)

    # ================= Титульный лист =================
    para(doc, 'АО «Казахстанский фонд гарантирования депозитов»',
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, bold=True)
    para(doc, 'Республика Казахстан, г. Алматы',
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=13)
    doc.add_paragraph()

    # Грифы согласования/утверждения
    grif = doc.add_table(rows=1, cols=2)
    grif.alignment = WD_TABLE_ALIGNMENT.CENTER
    grif.autofit = True
    left, right = grif.rows[0].cells
    for cell, label in ((left, 'СОГЛАСОВАНО'), (right, 'УТВЕРЖДАЮ')):
        cell.text = ''
        lines = [
            (label, True),
            ('Должность', False),
            ('_____________ /______________/', False),
            ('«____» ____________ 2026 г.', False),
        ]
        for text, bold in lines:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            _set_run(run, 13, bold=bold)

    for _ in range(4):
        doc.add_paragraph()

    para(doc, 'ИНФОРМАЦИОННАЯ СИСТЕМА', align=WD_ALIGN_PARAGRAPH.CENTER,
         indent=False, size=16, bold=True)
    para(doc, '«АВТОМАТИЗИРОВАННАЯ СИСТЕМА УЧЁТА» (ИС «АСУ»)',
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=18, bold=True)
    doc.add_paragraph()
    para(doc, 'ОПИСАНИЕ БАЗЫ ДАННЫХ', align=WD_ALIGN_PARAGRAPH.CENTER,
         indent=False, size=20, bold=True)
    doc.add_paragraph()
    para(doc, 'Разработано в соответствии с ГОСТ 34.321-96, ГОСТ 34.201-89, РД 50-34.698-90',
         align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=13)
    para(doc, 'Обозначение документа: АСУ.БД.ОП-01', align=WD_ALIGN_PARAGRAPH.CENTER,
         indent=False, size=13)
    para(doc, 'На правах рукописи', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, size=12, italic=True)

    for _ in range(7):
        doc.add_paragraph()
    para(doc, 'г. Алматы', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    para(doc, '2026', align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_page_break()

    # ================= Аннотация =================
    structural_heading(doc, 'АННОТАЦИЯ')
    para(doc, 'Настоящий документ содержит описание базы данных информационной системы '
              '«Автоматизированная система учёта» (далее — ИС «АСУ», Система) и разработан в составе '
              'рабочей документации на автоматизированную систему в соответствии с требованиями '
              'комплекса межгосударственных стандартов ГОСТ 34, действующих на территории '
              'Республики Казахстан.')
    para(doc, 'Описание базы данных выполнено на основе эталонной модели управления данными по '
              'ГОСТ 34.321-96 и включает: концептуальную схему информационной базы, описание '
              'сущностей и их атрибутов, внешние схемы (представления данных для категорий '
              'пользователей), внутреннюю схему (физическую реализацию в СУБД PostgreSQL), а также '
              'правила обеспечения целостности информационной базы.')
    para(doc, 'Документ предназначен для аналитиков, разработчиков, администраторов баз данных, '
              'специалистов по информационной безопасности и аудиторов.')
    doc.add_page_break()

    # ================= Содержание =================
    structural_heading(doc, 'СОДЕРЖАНИЕ')
    add_toc_field(doc)
    doc.add_page_break()

    # ================= 1 Общие положения =================
    heading(doc, '1 Общие положения')
    heading(doc, '1.1 Полное наименование системы', 2)
    para(doc, 'Полное наименование: информационная система «Автоматизированная система учёта» '
              'АО «Казахстанский фонд гарантирования депозитов». Условное обозначение: ИС «АСУ».')
    heading(doc, '1.2 Назначение документа', 2)
    para(doc, 'Документ определяет состав, структуру и правила ведения базы данных ИС «АСУ», '
              'обеспечивающей автоматизацию процессов складского учёта товарно-материальных запасов '
              '(ТМЗ), основных средств (ОС) и нематериальных активов (НМА), заявочного процесса, '
              'документооборота, уведомлений и интеграции с системой «1С».')
    heading(doc, '1.3 Нормативные ссылки', 2)
    para(doc, 'В настоящем документе использованы ссылки на следующие стандарты:')
    refs = [
        'ГОСТ 34.321-96 Информационные технологии. Система стандартов по базам данных. '
        'Эталонная модель управления данными;',
        'ГОСТ 34.201-89 Информационная технология. Комплекс стандартов на автоматизированные '
        'системы. Виды, комплектность и обозначение документов при создании автоматизированных систем;',
        'ГОСТ 34.601-90 Информационная технология. Комплекс стандартов на автоматизированные '
        'системы. Автоматизированные системы. Стадии создания;',
        'РД 50-34.698-90 Методические указания. Информационная технология. Комплекс стандартов и '
        'руководящих документов на автоматизированные системы. Требования к содержанию документов;',
        'СТ РК 34.015-2002 Информационная технология. Комплекс стандартов на автоматизированные '
        'системы. Техническое задание на создание автоматизированной системы.',
    ]
    for r in refs:
        para(doc, f'— {r}')
    heading(doc, '1.4 Основание для разработки', 2)
    para(doc, 'Основанием для разработки является техническое задание на создание ИС «АСУ» и '
              'решение о переводе процессов складского учёта и документооборота Фонда в '
              'автоматизированный контур.')
    heading(doc, '1.5 Среда реализации', 2)
    para(doc, 'База данных реализована в реляционной СУБД PostgreSQL версии 15. Доступ прикладного '
              'программного обеспечения к данным осуществляется средствами объектно-реляционного '
              'отображения Django ORM (Django 4.2). Изменения структуры выполняются исключительно '
              'механизмом миграций, что обеспечивает воспроизводимость схемы во всех средах '
              '(разработка, тестирование, промышленная эксплуатация).')

    # ================= 2 Термины =================
    heading(doc, '2 Термины, определения и сокращения')
    para(doc, 'В настоящем документе применены термины по ГОСТ 34.321-96, а также следующие '
              'термины с соответствующими определениями:')
    table_caption(doc, counters, 'Термины и определения')
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    cell_text(hdr[0], 'Термин', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(hdr[1], 'Определение', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(hdr[0], 'D9E2F3')
    shade_cell(hdr[1], 'D9E2F3')
    terms = [
        ('Информационная база', 'Совокупность данных, отражающих состояние предметной области и используемых в Системе (ГОСТ 34.321-96).'),
        ('Концептуальная схема', 'Непротиворечивое собрание предложений, выражающих необходимые высказывания о предметной области: сущности, атрибуты, связи и правила целостности (ГОСТ 34.321-96).'),
        ('Внешняя схема', 'Представление части информационной базы, предоставляемое конкретной категории пользователей (ГОСТ 34.321-96).'),
        ('Внутренняя схема', 'Описание физического представления данных в среде хранения СУБД (ГОСТ 34.321-96).'),
        ('Сущность', 'Объект предметной области, сведения о котором хранятся в информационной базе; реализуется таблицей БД.'),
        ('Атрибут', 'Свойство сущности; реализуется полем (столбцом) таблицы.'),
        ('ТМЗ', 'Товарно-материальные запасы.'),
        ('ОС', 'Основные средства.'),
        ('НМА', 'Нематериальные активы.'),
        ('МОЛ', 'Материально ответственное лицо.'),
        ('АХС', 'Административно-хозяйственная служба.'),
        ('PK / FK / UK', 'Первичный ключ / внешний ключ / уникальный ключ (ограничение уникальности).'),
        ('CASCADE / PROTECT / SET NULL', 'Стратегии ссылочной целостности при удалении родительской записи.'),
        ('GenericFK (полиморфная связь)', 'Ссылка на запись произвольной таблицы через пару полей «тип содержимого + идентификатор» (механизм ContentType).'),
        ('M2M', 'Связь «многие ко многим», реализуемая промежуточной таблицей.'),
    ]
    for term, definition in terms:
        row = t.add_row().cells
        cell_text(row[0], term, 12)
        cell_text(row[1], definition, 12)
    doc.add_paragraph()

    # ================= 3 Эталонная модель =================
    heading(doc, '3 Уровни описания данных по эталонной модели ГОСТ 34.321-96')
    para(doc, 'В соответствии с эталонной моделью управления данными описание базы данных ИС «АСУ» '
              'ведётся на трёх уровнях:')
    para(doc, '— концептуальный уровень (раздел 4) — сущности предметной области, их атрибуты, связи '
              'и правила целостности, независимые от способа физического хранения;')
    para(doc, '— внешний уровень (раздел 5) — представления информационной базы для категорий '
              'пользователей Системы в соответствии с моделью разграничения доступа;')
    para(doc, '— внутренний уровень (раздел 6) — физическая реализация в СУБД PostgreSQL: таблицы, '
              'типы данных, индексы, ограничения.')
    para(doc, 'Предметная область Системы декомпозирована на семь функциональных контуров: '
              'организационная структура и доступы; нормативно-справочная информация; номенклатура '
              'активов; складской учёт; заявочный процесс; документооборот; уведомления и интеграции. '
              'Ключевой принцип концептуальной схемы: справочники являются источником мастер-данных, '
              'заявки и документы фиксируют бизнес-процесс и его юридическую значимость, а складские '
              'движения ведут журнал фактических операций.')
    add_image(doc, images['overview'])
    figure_caption(doc, counters, 'Функциональные контуры информационной базы ИС «АСУ»')

    # ================= 4 Концептуальная схема =================
    heading(doc, '4 Концептуальная схема информационной базы')

    section_intro = {
        'Организационная структура и права доступа':
            ('Контур определяет состав пользователей Системы и их полномочия. Модель доступа '
             'трёхуровневая: базовая роль хранится в атрибуте role сущности users_user; типовые '
             'права по должностям задаются правилами users_positionaccessrule; индивидуальные '
             'исключения — записями users_useraccessoverride, имеющими наивысший приоритет. '
             'Подразделения образуют иерархию через ссылку parent_id, атрибут supervisor_id '
             'обеспечивает маршрутизацию согласования заявок на непосредственного руководителя.',
             'org', 'Схема данных контура организационной структуры и прав доступа'),
        'Нормативно-справочная информация':
            ('Справочники являются источником мастер-данных для всех операционных контуров. '
             'Карточка актива references_asset универсальна для всех типов активов и содержит '
             'атрибуты интеграции с системой «1С» (source_1c_id, last_sync_at). Категории и группы '
             'реализованы единой сущностью references_assetcategory с иерархией через parent_id.',
             'refs', 'Схема данных контура нормативно-справочной информации'),
        'Складской учёт':
            ('Контур разделяет текущее состояние остатков (assets_warehousestock) и историю операций '
             '(assets_stockmovement). Журнал движений связывается с документом-основанием '
             'полиморфной ссылкой (document_type_id + document_id). Правила контроля критических '
             'остатков и их срабатывания обеспечивают своевременное уведомление ответственных лиц.',
             'stock', 'Схема данных контура складского учёта'),
        'Заявочный процесс':
            ('Заявка фиксирует бизнес-намерение сотрудника. Позиции заявки могут ссылаться как на '
             'конкретный актив, так и на группу номенклатуры. Журнал согласования отделён от заявки: '
             'атрибут status отражает текущее состояние, а requests_requestapproval хранит полную '
             'историю действий. Маршрут согласования настраивается по каждому виду заявки.',
             'requests', 'Схема данных контура заявочного процесса'),
        'Документооборот':
            ('Документы построены по схеме «заголовок — позиции» и наследуют общий набор атрибутов '
             'абстрактной сущности BaseDocument: номер (присваивается после финального подписания), '
             'дата, статус, автор, метки времени. Подписи вынесены в единую сущность '
             'documents_documentsignature с полиморфной связью — одна таблица обслуживает все пять '
             'типов документов.',
             'documents', 'Схема данных контура документооборота'),
        'Уведомления и интеграции':
            ('Уведомления используют полиморфную связь с любым объектом Системы. Отправка электронной '
             'почты журналируется с фиксацией статуса и текста ошибки. Каждый сеанс обмена с системой '
             '«1С» регистрируется в журнале integrations_synclog.',
             'notifications', 'Схема данных контура уведомлений и интеграций'),
    }

    sub = 1
    for section_name, tables in ALL_SECTIONS:
        heading(doc, f'4.{sub} {section_name}', 2)
        intro, img_key, fig_title = section_intro[section_name]
        para(doc, intro)
        add_image(doc, images[img_key])
        figure_caption(doc, counters, fig_title)
        for tname, purpose, fields in tables:
            para(doc, f'Сущность {tname}. {purpose}')
            table_caption(doc, counters, f'Атрибуты сущности {tname}')
            fields_table(doc, fields)
            doc.add_paragraph()
        sub += 1

    # Статусные модели
    heading(doc, f'4.{sub} Статусные модели процессных сущностей', 2)
    para(doc, 'Заявки и документы являются процессными сущностями: атрибут status определяет '
              'положение объекта в жизненном цикле. Переходы состояний выполняются прикладным '
              'программным обеспечением и фиксируются в журналах согласований и подписей.')
    add_image(doc, images['statuses'])
    figure_caption(doc, counters, 'Статусные модели заявок, документов и закреплений активов')

    # ================= 5 Внешние схемы =================
    heading(doc, '5 Внешние схемы (представления для категорий пользователей)')
    para(doc, 'Внешние схемы реализуются программно на уровне сервисов доступа и определяют, какая '
              'часть информационной базы доступна каждой категории пользователей:')
    table_caption(doc, counters, 'Внешние схемы информационной базы')
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    cell_text(hdr[0], 'Категория пользователей', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(hdr[1], 'Доступная часть информационной базы', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(hdr[0], 'D9E2F3')
    shade_cell(hdr[1], 'D9E2F3')
    ext = [
        ('Сотрудник (USER)', 'Собственный профиль; собственные заявки и заявки своего подразделения; уведомления; закреплённые активы.'),
        ('Руководитель подразделения (DEPT_HEAD)', 'Дополнительно: заявки подчинённых на этапе согласования; данные своего подразделения.'),
        ('Работник и руководитель АХС (AHS_WORKER, AHS_HEAD)', 'Заявки на согласовании и исполнении; документооборот; справочники; отчётность.'),
        ('МОЛ (MOL_WAREHOUSE, MOL_NMA)', 'Складские остатки, движения, закрепления; загрузка остатков; приходные документы; алармы остатков.'),
        ('Администратор (ADMIN)', 'Полный доступ ко всем сущностям, управление пользователями, правами, синхронизацией с «1С».'),
    ]
    for role, access in ext:
        row = t.add_row().cells
        cell_text(row[0], role, 12)
        cell_text(row[1], access, 12)
    doc.add_paragraph()
    para(doc, 'Итоговый набор полномочий пользователя вычисляется как суперпозиция: роль → правила '
              'по должности → индивидуальные исключения. Матрица полномочий администрируется в '
              'подсистеме управления доступом и может выгружаться для аудита.')

    # ================= 6 Внутренняя схема =================
    heading(doc, '6 Внутренняя схема (физическая реализация)')
    heading(doc, '6.1 Общие сведения', 2)
    para(doc, 'Все сущности концептуальной схемы реализованы таблицами PostgreSQL 15 в схеме public. '
              'Первичные ключи — суррогатные (bigint, автоинкремент). Денежные и количественные '
              'атрибуты хранятся в типе numeric (decimal) с фиксированной точностью: количество — '
              'numeric(12,2), стоимость — numeric(15,2). Метки времени — timestamp with time zone.')
    heading(doc, '6.2 Ограничения целостности уровня СУБД', 2)
    para(doc, 'В базе данных применяются следующие декларативные ограничения:')
    para(doc, '— ограничения уникальности: users_department.code, references_counterparty.bin, '
              'references_asset.code, references_asset.source_1c_id, requests_assetrequest.number, '
              'составные ключи (normalized_position, permission_code), (user_id, permission_code), '
              '(rule_id, stock_id), (request_type_id, order);')
    para(doc, '— ограничения ссылочной целостности (внешние ключи) со стратегиями PROTECT, CASCADE, '
              'SET NULL согласно таблице ниже;')
    para(doc, '— ограничения NOT NULL на обязательные атрибуты сущностей.')
    table_caption(doc, counters, 'Стратегии ссылочной целостности')
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, name in enumerate(['Стратегия', 'Область применения', 'Обоснование']):
        cell_text(hdr[i], name, 12, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[i], 'D9E2F3')
    integrity = [
        ('PROTECT', 'Активы в позициях документов и заявок; категории; виды заявок; контрагенты и склады в накладных; единицы измерения.',
         'Справочная запись не может быть удалена, пока используется в юридически значимых данных.'),
        ('CASCADE', 'Позиции документов и заявок; договоры контрагентов; индивидуальные права; срабатывания алармов.',
         'Дочерние записи не имеют самостоятельного смысла и удаляются вместе с родительской.'),
        ('SET NULL', 'Руководители; МОЛ; исполнители операций; склады и подразделения в исторических записях.',
         'Историческая запись сохраняется при удалении связанного объекта.'),
    ]
    for strat, where, sense in integrity:
        row = t.add_row().cells
        cell_text(row[0], strat, 12)
        cell_text(row[1], where, 12)
        cell_text(row[2], sense, 12)
    doc.add_paragraph()
    heading(doc, '6.3 Индексы', 2)
    para(doc, 'Помимо индексов, автоматически создаваемых для первичных, уникальных и внешних '
              'ключей, для обеспечения производительности рекомендованы следующие индексы:')
    table_caption(doc, counters, 'Рекомендуемые индексы')
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    cell_text(hdr[0], 'Таблица', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(hdr[1], 'Индексируемые атрибуты', 12, True, WD_ALIGN_PARAGRAPH.CENTER)
    shade_cell(hdr[0], 'D9E2F3')
    shade_cell(hdr[1], 'D9E2F3')
    idx = [
        ('requests_assetrequest', 'status; initiator_id; created_at; request_type_id'),
        ('requests_requestapproval', 'request_id; approver_id; created_at'),
        ('documents_* (заголовки)', 'status; created_by_id; created_at; date'),
        ('notifications_notification', 'составной (recipient_id, is_read, created_at)'),
        ('assets_stockmovement', 'asset_id; warehouse_id; performed_at; movement_type'),
        ('references_asset', 'asset_type; category_id; group_id; source_1c_id'),
    ]
    for tab, i_ in idx:
        row = t.add_row().cells
        cell_text(row[0], tab, 12)
        cell_text(row[1], i_, 12)
    doc.add_paragraph()
    heading(doc, '6.4 Хранение файлов', 2)
    para(doc, 'Файловые данные (фотографии пользователей, PDF-файлы договоров) хранятся в файловом '
              'хранилище; в базе данных сохраняются относительные пути (users/photos/, '
              'contracts/pdfs/).')

    # ================= 7 Правила ведения =================
    heading(doc, '7 Правила ведения информационной базы')
    heading(doc, '7.1 Наполнение и актуализация', 2)
    para(doc, 'Справочники ведутся уполномоченными сотрудниками АХС и администраторами через '
              'пользовательский интерфейс Системы. Номенклатура и остатки могут загружаться из '
              'системы «1С»; ключом сопоставления служит атрибут source_1c_id (резервный ключ — '
              'code). Каждая загрузка регистрируется в журнале integrations_synclog с фиксацией '
              'количества созданных и обновлённых записей; импорт является идемпотентным.')
    heading(doc, '7.2 Аудит и юридическая значимость', 2)
    para(doc, 'Аудиторскую функцию выполняют журналы: складских движений (assets_stockmovement), '
              'согласований заявок (requests_requestapproval), подписей документов '
              '(documents_documentsignature), электронной почты (notifications_emaillog) и '
              'синхронизаций (integrations_synclog). Номера заявкам и документам присваиваются '
              'автоматически в формате NNN/ГГГГ; номер документа присваивается только после '
              'финального подписания.')
    heading(doc, '7.3 Резервное копирование и восстановление', 2)
    para(doc, 'Резервное копирование базы данных выполняется штатными средствами PostgreSQL '
              '(pg_dump / непрерывное архивирование WAL) по регламенту эксплуатирующего '
              'подразделения. Восстановление проверяется на тестовой среде.')

    # ================= 8 Ограничения и развитие =================
    heading(doc, '8 Ограничения текущей реализации и направления развития')
    limits = [
        'Сущность assets_warehousestock хранит одну строку остатка на актив (связь «один к одному»). '
        'Для мультискладского учёта необходимо перейти к связи «многие к одному» с уникальным '
        'ключом (asset_id, warehouse_id).',
        'Сущность references_assetcategory используется одновременно как категория и как группа; '
        'при росте иерархии целесообразно ввести явный признак уровня.',
        'Атрибут number документов не имеет ограничения уникальности на уровне СУБД; нумерация '
        'ведётся программно в разрезе типа документа.',
        'Для финансовых атрибутов рекомендуется добавить контрольные ограничения СУБД: '
        'quantity > 0, unit_price >= 0, total >= 0.',
        'После перехода к мультискладскому учёту ключом сопоставления с «1С» должна стать пара '
        '(source_1c_id, warehouse_code).',
    ]
    for i, l in enumerate(limits, 1):
        para(doc, f'{i}) {l}')

    doc.add_page_break()

    # ================= Лист регистрации изменений =================
    structural_heading(doc, 'ЛИСТ РЕГИСТРАЦИИ ИЗМЕНЕНИЙ')
    t = doc.add_table(rows=2, cols=9)
    t.style = 'Table Grid'
    headers = ['Изм.', 'Номера листов (страниц)', '', '', '', 'Всего листов', '№ документа',
               'Подпись', 'Дата']
    sub_headers = ['', 'изменённых', 'заменённых', 'новых', 'аннулированных', '', '', '', '']
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, 10, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(t.rows[0].cells[i], 'D9E2F3')
    for i, h in enumerate(sub_headers):
        cell_text(t.rows[1].cells[i], h, 10, True, WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(t.rows[1].cells[i], 'D9E2F3')
    for _ in range(8):
        row = t.add_row()
        for cell in row.cells:
            cell_text(cell, '', 10)

    out = os.path.join(BASE_DIR, 'ASU_Описание_БД_ГОСТ_34.321-96.docx')
    doc.save(out)
    return out


def main():
    images = {
        'overview': diagram_overview(),
        'org': diagram_org(),
        'refs': diagram_refs(),
        'stock': diagram_stock(),
        'requests': diagram_requests(),
        'documents': diagram_documents(),
        'notifications': diagram_notifications(),
        'statuses': diagram_statuses(),
    }
    out = build(images)
    print('OK:', out)


if __name__ == '__main__':
    main()
