# -*- coding: utf-8 -*-
"""Генерация docx-описания базы данных ИС «АСУ» со схемами."""

import os

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from db_doc_data import ALL_SECTIONS

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(BASE_DIR, '_db_doc_images')
os.makedirs(IMG_DIR, exist_ok=True)

plt.rcParams['font.family'] = 'DejaVu Sans'

HEADER_COLOR = '#1F3864'
BODY_COLOR = '#EAF0F8'
EDGE_COLOR = '#2F5597'
ACCENT = {
    'org': '#1F3864',
    'ref': '#2E7D32',
    'stock': '#B26A00',
    'req': '#7B1FA2',
    'doc': '#C62828',
    'notif': '#00695C',
}


def draw_entity(ax, x, y, w, title, fields, color):
    """Рисует сущность: заголовок + поля. Возвращает (cx, cy, w, h)."""
    line_h = 0.34
    h = 0.52 + line_h * len(fields) + 0.14
    box = FancyBboxPatch((x, y - h), w, h, boxstyle='round,pad=0.03,rounding_size=0.08',
                         linewidth=1.4, edgecolor=color, facecolor=BODY_COLOR, zorder=2)
    ax.add_patch(box)
    head = FancyBboxPatch((x, y - 0.52), w, 0.52, boxstyle='round,pad=0.02,rounding_size=0.08',
                          linewidth=0, facecolor=color, zorder=3)
    ax.add_patch(head)
    ax.text(x + w / 2, y - 0.27, title, ha='center', va='center', fontsize=8.5,
            color='white', fontweight='bold', zorder=4)
    for i, f in enumerate(fields):
        ax.text(x + 0.12, y - 0.66 - line_h * i - line_h / 2, f, ha='left', va='center',
                fontsize=7.2, color='#1a1a2e', zorder=4)
    return (x + w / 2, y - h / 2, w, h)


def connect(ax, a, b, label='', color='#555577', rad=0.08):
    """Стрелка между центрами сущностей a и b."""
    ax1, ay1 = a[0], a[1]
    bx1, by1 = b[0], b[1]
    arrow = FancyArrowPatch((ax1, ay1), (bx1, by1),
                            connectionstyle=f'arc3,rad={rad}',
                            arrowstyle='-|>', mutation_scale=11,
                            linewidth=1.1, color=color, zorder=1, alpha=0.85)
    ax.add_patch(arrow)
    if label:
        mx, my = (ax1 + bx1) / 2, (ay1 + by1) / 2
        ax.text(mx, my + 0.14, label, ha='center', va='center', fontsize=6.6,
                color=color, zorder=5,
                bbox=dict(boxstyle='round,pad=0.15', fc='white', ec='none', alpha=0.85))


def save_fig(fig, name):
    path = os.path.join(IMG_DIR, name)
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# Схема 1. Обзор доменных контуров
# ============================================================

def diagram_overview():
    fig, ax = plt.subplots(figsize=(11, 6.2))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')

    blocks = [
        (0.4, 7.6, 'Оргструктура и права', ['users_department', 'users_user', 'references_position', 'users_positionaccessrule', 'users_useraccessoverride'], ACCENT['org']),
        (5.2, 7.6, 'Справочники (НСИ)', ['references_counterparty', 'references_contract', 'references_requesttype', 'references_assetcategory', 'references_unitofmeasure', 'references_warehouse', 'references_limitnorm'], ACCENT['ref']),
        (10.0, 7.6, 'Номенклатура', ['references_asset'], ACCENT['ref']),
        (0.4, 3.6, 'Заявочный процесс', ['requests_assetrequest', 'requests_assetrequestitem', 'requests_requestapproval', 'requests_approvalstep'], ACCENT['req']),
        (5.2, 3.9, 'Складской учёт', ['assets_warehousestock', 'assets_stockmovement', 'assets_assetassignment', 'assets_stockalertrule', 'assets_stockalertstate'], ACCENT['stock']),
        (10.0, 5.6, 'Документооборот', ['5 документов «шапка + позиции»', 'documents_documentsignature', 'documents_commissionmember'], ACCENT['doc']),
        (10.0, 2.4, 'Уведомления и интеграции', ['notifications_notification', 'notifications_emaillog', 'integrations_synclog'], ACCENT['notif']),
    ]
    pos = {}
    for x, y, title, fields, color in blocks:
        pos[title] = draw_entity(ax, x, y, 3.7, title, fields, color)

    connect(ax, pos['Оргструктура и права'], pos['Заявочный процесс'], 'инициатор, согласующие')
    connect(ax, pos['Справочники (НСИ)'], pos['Номенклатура'], 'категории, ЕИ')
    connect(ax, pos['Номенклатура'], pos['Складской учёт'], 'остатки, движения', rad=-0.15)
    connect(ax, pos['Заявочный процесс'], pos['Складской учёт'], 'выдача')
    connect(ax, pos['Документооборот'], pos['Складской учёт'], 'основание движения')
    connect(ax, pos['Номенклатура'], pos['Документооборот'], 'позиции')
    connect(ax, pos['Документооборот'], pos['Уведомления и интеграции'], 'события')
    ax.set_title('Доменные контуры базы данных ИС «АСУ»', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '01_overview.png')


# ============================================================
# Схема 2. Организационная структура и права
# ============================================================

def diagram_org():
    fig, ax = plt.subplots(figsize=(10.5, 6))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 8)
    ax.axis('off')
    c = ACCENT['org']

    dep = draw_entity(ax, 0.5, 7.4, 3.4, 'users_department', ['id PK', 'code UK', 'name', 'head_id FK', 'parent_id FK (self)'], c)
    usr = draw_entity(ax, 4.8, 7.8, 3.6, 'users_user', ['id PK', 'username UK', 'ФИО, фото, телефон', 'role (choices)', 'position_ref_id FK', 'department_id FK', 'supervisor_id FK (self)'], c)
    posn = draw_entity(ax, 9.3, 7.4, 3.2, 'references_position', ['id PK', 'name UK', 'code UK', 'is_active'], ACCENT['ref'])
    rule = draw_entity(ax, 0.9, 3.0, 4.0, 'users_positionaccessrule', ['id PK', 'normalized_position (idx)', 'permission_code', 'is_allowed', 'UK(position+code)'], c)
    ovr = draw_entity(ax, 7.4, 3.0, 4.0, 'users_useraccessoverride', ['id PK', 'user_id FK', 'permission_code', 'mode: GRANT/DENY', 'UK(user+code)'], c)

    connect(ax, dep, dep, '', rad=0.6)
    connect(ax, usr, dep, 'department 1—N', rad=-0.1)
    connect(ax, dep, usr, 'head', rad=0.25)
    connect(ax, usr, posn, 'position_ref N—1')
    connect(ax, usr, usr, 'supervisor (self)', rad=0.6)
    connect(ax, ovr, usr, 'user N—1', rad=-0.15)
    connect(ax, rule, posn, 'по названию должности', color='#999999', rad=0.2)
    ax.set_title('Контур 1. Организационная структура и права доступа', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '02_org.png')


# ============================================================
# Схема 3. Справочники и номенклатура
# ============================================================

def diagram_refs():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8.6)
    ax.axis('off')
    c = ACCENT['ref']

    cp = draw_entity(ax, 0.4, 8.2, 3.2, 'references_counterparty', ['id PK', 'bin UK', 'name', 'is_active'], c)
    ctr = draw_entity(ax, 0.4, 4.6, 3.2, 'references_contract', ['id PK', 'counterparty_id FK', 'contract_date', 'valid_until', 'pdf_file'], c)
    cat = draw_entity(ax, 4.2, 8.2, 3.4, 'references_assetcategory', ['id PK', 'code UK', 'asset_type', 'parent_id FK (self)'], c)
    uom = draw_entity(ax, 4.2, 4.9, 3.0, 'references_unitofmeasure', ['id PK', 'name UK', 'code UK'], c)
    wh = draw_entity(ax, 4.2, 2.6, 3.0, 'references_warehouse', ['id PK', 'code UK', 'department_id FK'], c)
    asset = draw_entity(ax, 8.4, 7.6, 3.8, 'references_asset', ['id PK', 'code UK', 'asset_type', 'category_id FK PROTECT', 'group_id FK', 'unit_of_measure_ref_id FK', 'unit_price', 'inventory_number', 'source_1c_id UK'], c)
    rt = draw_entity(ax, 8.6, 2.6, 3.4, 'references_requesttype', ['id PK', 'code UK', 'asset_type', 'requires_long_term_use'], c)
    lim = draw_entity(ax, 0.4, 2.2, 3.2, 'references_limitnorm', ['id PK', 'asset_type', 'quantity_limit', 'department_id FK'], c)

    connect(ax, ctr, cp, 'N—1 CASCADE')
    connect(ax, cat, cat, 'parent (self)', rad=0.6)
    connect(ax, asset, cat, 'category / group', rad=-0.1)
    connect(ax, asset, uom, 'unit_of_measure_ref', rad=0.1)
    ax.set_title('Контур 2. Справочники и номенклатура', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '03_refs.png')


# ============================================================
# Схема 4. Складской учёт
# ============================================================

def diagram_stock():
    fig, ax = plt.subplots(figsize=(11, 6.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8.6)
    ax.axis('off')
    c = ACCENT['stock']

    asset = draw_entity(ax, 0.4, 8.2, 3.2, 'references_asset', ['id PK', 'code UK', 'unit_price'], ACCENT['ref'])
    stock = draw_entity(ax, 4.6, 8.2, 3.6, 'assets_warehousestock', ['id PK', 'asset_id 1—1 UK', 'warehouse_id FK', 'quantity', 'total_amount', 'balance_date'], c)
    mov = draw_entity(ax, 9.4, 8.2, 4.0, 'assets_stockmovement', ['id PK', 'asset_id FK', 'movement_type', 'quantity, unit_price', 'from/to/performed_by FK', 'document_type + document_id', '(GenericFK на документ)'], c)
    asg = draw_entity(ax, 0.4, 4.2, 3.6, 'assets_assetassignment', ['id PK', 'asset_id FK', 'user_id FK', 'warehouse_id FK', 'status'], c)
    rule = draw_entity(ax, 4.8, 3.9, 3.6, 'assets_stockalertrule', ['id PK', 'threshold_quantity', 'M2M: recipients, groups,', 'assets, warehouses', 'message_template'], c)
    state = draw_entity(ax, 9.6, 4.0, 3.6, 'assets_stockalertstate', ['id PK', 'rule_id FK', 'stock_id FK', 'is_active', 'UK(rule+stock)'], c)

    connect(ax, stock, asset, '1—1')
    connect(ax, mov, asset, 'N—1', rad=0.2)
    connect(ax, asg, asset, 'N—1', rad=-0.1)
    connect(ax, state, rule, 'N—1')
    connect(ax, state, stock, 'N—1', rad=-0.2)
    ax.set_title('Контур 3. Складской учёт и алармы остатков', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '04_stock.png')


# ============================================================
# Схема 5. Заявочный процесс
# ============================================================

def diagram_requests():
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 8)
    ax.axis('off')
    c = ACCENT['req']

    rt = draw_entity(ax, 0.4, 7.6, 3.4, 'references_requesttype', ['id PK', 'code UK', 'asset_type'], ACCENT['ref'])
    req = draw_entity(ax, 5.0, 7.8, 3.8, 'requests_assetrequest', ['id PK', 'number UK (NNN/ГГГГ)', 'request_type_id FK', 'status', 'initiator_id FK', 'M2M issue_responsibles'], c)
    step = draw_entity(ax, 0.4, 3.4, 3.6, 'requests_approvalstep', ['id PK', 'request_type_id FK', 'order UK', 'approver_role', 'requires_supervisor'], c)
    item = draw_entity(ax, 5.2, 3.2, 3.6, 'requests_assetrequestitem', ['id PK', 'request_id FK', 'requested_group_id FK', 'asset_id / issued_asset_id FK', 'quantity_requested/issued'], c)
    appr = draw_entity(ax, 10.0, 6.0, 3.4, 'requests_requestapproval', ['id PK', 'request_id FK', 'approver_id FK', 'action, signed_at'], c)

    connect(ax, req, rt, 'N—1 PROTECT')
    connect(ax, step, rt, 'маршрут N—1', rad=-0.1)
    connect(ax, item, req, 'N—1 CASCADE')
    connect(ax, appr, req, 'журнал N—1', rad=-0.15)
    ax.set_title('Контур 4. Заявочный процесс', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '05_requests.png')


# ============================================================
# Схема 6. Документооборот
# ============================================================

def diagram_documents():
    fig, ax = plt.subplots(figsize=(11.5, 7))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 9.4)
    ax.axis('off')
    c = ACCENT['doc']

    inv = draw_entity(ax, 0.3, 9.0, 3.3, 'incominginvoice', ['counterparty_id FK', 'warehouse_id FK', 'mol_warehouse_id FK', '+ items (позиции)'], c)
    act = draw_entity(ax, 3.9, 9.0, 3.3, 'writeoffact', ['act_type', 'total_amount', 'is_representative', '+ items (позиции)'], c)
    pet = draw_entity(ax, 7.5, 9.0, 3.3, 'petition', ['legal_basis', '+ items (позиции)'], c)
    prot = draw_entity(ax, 11.1, 9.0, 3.5, 'commissionprotocol', ['petition_id FK', 'decision_text', '+ attachment_items'], c)
    trans = draw_entity(ax, 0.3, 5.6, 3.3, 'internaltransferinvoice', ['from_user_id FK', 'to_user_id FK', 'asset_type: OS/NMA', '+ items (позиции)'], c)
    base = draw_entity(ax, 4.4, 5.2, 3.6, 'BaseDocument (абстракт)', ['number, date', 'status (workflow)', 'created_by_id FK', 'created_at, updated_at'], '#555566')
    sig = draw_entity(ax, 8.3, 5.2, 3.5, 'documentsignature', ['GenericFK: document_type', '+ document_id', 'signer_id FK', 'signed_at, role_label'], c)
    comm = draw_entity(ax, 12.1, 5.0, 2.8, 'commissionmember', ['user_id FK', 'act / petition /', 'protocol FK'], c)

    for e in (inv, act, pet, prot, trans):
        connect(ax, e, base, '', color='#888899', rad=0.05)
    connect(ax, prot, pet, 'основание', rad=0.2)
    connect(ax, sig, base, 'подписи всех типов', rad=-0.1)
    connect(ax, comm, act, '', rad=0.25)
    ax.set_title('Контур 5. Документооборот («шапка + позиции», единые подписи)', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '06_documents.png')


# ============================================================
# Схема 7. Уведомления и интеграции
# ============================================================

def diagram_notifications():
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 6.4)
    ax.axis('off')
    c = ACCENT['notif']

    usr = draw_entity(ax, 0.4, 6.0, 3.0, 'users_user', ['id PK'], ACCENT['org'])
    ntf = draw_entity(ax, 4.4, 6.0, 3.8, 'notifications_notification', ['id PK', 'recipient_id FK', 'notification_type', 'GenericFK: related_content_type', '+ related_object_id', 'is_read'], c)
    eml = draw_entity(ax, 9.2, 6.0, 3.4, 'notifications_emaillog', ['id PK', 'recipient_email', 'status: SENT/FAILED', 'related_notification_id FK'], c)
    sync = draw_entity(ax, 4.6, 2.2, 3.6, 'integrations_synclog', ['id PK', 'sync_type, status', 'created/updated_count', 'is_stub'], c)

    connect(ax, ntf, usr, 'N—1')
    connect(ax, eml, ntf, 'N—1 SET_NULL')
    ax.set_title('Контур 6. Уведомления и интеграции', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '07_notifications.png')


# ============================================================
# Схема 8. Статусные модели
# ============================================================

def diagram_statuses():
    fig, ax = plt.subplots(figsize=(11, 5.6))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 7.2)
    ax.axis('off')

    def flow(y, states, color, title):
        ax.text(0.3, y + 0.75, title, fontsize=10, fontweight='bold', color=color)
        boxes = []
        x = 0.3
        for s in states:
            w = 0.28 + 0.115 * len(s)
            box = FancyBboxPatch((x, y - 0.3), w, 0.6, boxstyle='round,pad=0.03,rounding_size=0.12',
                                 linewidth=1.2, edgecolor=color, facecolor='#F4F6FB')
            ax.add_patch(box)
            ax.text(x + w / 2, y, s, ha='center', va='center', fontsize=7.4, color='#1a1a2e')
            boxes.append((x, w))
            x += w + 0.55
        for i in range(len(boxes) - 1):
            x0 = boxes[i][0] + boxes[i][1]
            x1 = boxes[i + 1][0]
            ax.annotate('', xy=(x1, y), xytext=(x0 + 0.05, y),
                        arrowprops=dict(arrowstyle='-|>', color=color, lw=1.2))

    flow(5.9, ['DRAFT', 'PENDING_SUPERVISOR', 'APPROVED_SUPERVISOR', 'APPROVED_AHS_HEAD', 'APPROVED', 'EXECUTED'],
         ACCENT['req'], 'Жизненный цикл заявки (requests_assetrequest.status)')
    ax.text(0.3, 4.9, 'Альтернативные состояния: SENT_FOR_REVISION, REJECTED, CANCELLED',
            fontsize=8.2, color='#555566')

    flow(3.2, ['DRAFT', 'PENDING_AHS_APPROVAL', 'PENDING_SIGNATURE', 'PARTIALLY_SIGNED', 'SIGNED'],
         ACCENT['doc'], 'Жизненный цикл документа (documents_*.status)')
    ax.text(0.3, 2.2, 'Дополнительные состояния: PENDING_CHANGE_APPROVAL, SENT_FOR_REVISION, REJECTED, CANCELLED',
            fontsize=8.2, color='#555566')

    flow(1.0, ['ACTIVE', 'RETURNED / WRITTEN_OFF'], ACCENT['stock'],
         'Статусы закрепления (assets_assetassignment.status)')
    ax.set_title('Статусные модели ключевых процессов', fontsize=12, fontweight='bold', pad=14)
    return save_fig(fig, '08_statuses.png')


# ============================================================
# DOCX
# ============================================================

def set_font(run, size=14, bold=False, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text, size=14, bold=False, align=None, space_after=6, color=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size, bold, color, italic)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 16, 3: 14}
    colors = {1: '1F3864', 2: '2F5597', 3: '2F5597'}
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, sizes[level], bold=True, color=colors[level])
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_fields_table(doc, fields):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, name in enumerate(['Поле', 'Тип / связь', 'Назначение']):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(name)
        set_font(run, 12, bold=True, color='FFFFFF')
        shade_cell(hdr[i], '1F3864')
    for fname, ftype, fdesc in fields:
        row = table.add_row().cells
        for i, val in enumerate([fname, ftype, fdesc]):
            run = row[i].paragraphs[0].add_run(val)
            set_font(run, 11 if i < 2 else 11)
    widths = [Cm(4.6), Cm(5.6), Cm(6.8)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w
    return table


def add_image(doc, path, width_cm=16.5, caption=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        set_font(run, 11, italic=True, color='555566')
        cap.paragraph_format.space_after = Pt(12)


def build_document(images):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    # ---------- Титульный лист ----------
    for _ in range(6):
        doc.add_paragraph()
    add_para(doc, 'АО «Казахстанский фонд гарантирования депозитов»', 16, True, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, 'ИНФОРМАЦИОННАЯ СИСТЕМА «АСУ»', 22, True, WD_ALIGN_PARAGRAPH.CENTER, color='1F3864')
    add_para(doc, 'Описание архитектуры и структуры базы данных', 18, True, WD_ALIGN_PARAGRAPH.CENTER, color='2F5597')
    doc.add_paragraph()
    add_para(doc, 'Складской учёт ТМЗ, ОС и НМА · Заявки · Документооборот · Уведомления · Интеграция с 1С',
             13, False, WD_ALIGN_PARAGRAPH.CENTER, color='555566')
    for _ in range(8):
        doc.add_paragraph()
    add_para(doc, 'Технологический стек: Django 4.2 (ORM) · PostgreSQL 15', 13, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Версия документа: 1.0 · Июль 2026', 13, False, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---------- 1. Общие сведения ----------
    add_heading(doc, '1. Общие сведения', 1)
    add_para(doc, 'Настоящий документ содержит полное описание базы данных информационной системы «АСУ» — '
                  'корпоративного портала складского учёта товарно-материальных запасов (ТМЗ), основных средств (ОС) '
                  'и нематериальных активов (НМА). Документ предназначен для аналитиков, разработчиков, администраторов '
                  'баз данных и аудиторов и позволяет получить исчерпывающее представление о структуре данных, связях '
                  'между таблицами и принятых архитектурных решениях.')
    add_para(doc, 'База данных реализована на PostgreSQL 15 и управляется через Django ORM. Все таблицы создаются '
                  'и изменяются исключительно через механизм миграций, что обеспечивает воспроизводимость схемы в любом '
                  'окружении. Аутентификация построена на пользовательской модели users_user, расширяющей стандартную '
                  'модель Django.')
    add_para(doc, 'Ключевая архитектурная идея системы: справочники являются источником мастер-данных, заявки и '
                  'документы фиксируют бизнес-процесс и его юридическую значимость, а складские движения ведут журнал '
                  'фактических операций. Такое разделение позволяет независимо отследить намерение сотрудника, '
                  'оформленный документ и фактическое изменение остатков.')

    # ---------- 2. Архитектурный обзор ----------
    add_heading(doc, '2. Архитектурный обзор', 1)
    add_para(doc, 'Логически схема базы данных делится на семь взаимосвязанных контуров:')
    contours = [
        'Организационная структура и доступы — пользователи, подразделения, должности, права по должностям и индивидуальные права.',
        'Нормативно-справочная информация — контрагенты, договоры, виды заявок, категории и группы, единицы измерения, склады, лимиты.',
        'Номенклатура — единая карточка актива references_asset для ТМЗ, ОС, НМА и представительских ТМЗ.',
        'Складской учёт — текущие остатки, журнал движений, закрепления за сотрудниками, алармы критических остатков.',
        'Заявочный процесс — заявка, позиции, настраиваемые этапы согласования, журнал действий.',
        'Документооборот — приходные накладные, акты списания, ходатайства, протоколы комиссии, внутренние перемещения, универсальные подписи.',
        'Уведомления и интеграции — in-app уведомления, журнал email, журнал синхронизаций с 1С.',
    ]
    for i, c in enumerate(contours, 1):
        add_para(doc, f'{i}. {c}', space_after=3)
    doc.add_paragraph()
    add_image(doc, images['overview'], 16.5, 'Рисунок 1. Доменные контуры базы данных и основные потоки данных между ними')

    # ---------- 3+. Контуры с таблицами ----------
    section_intro = {
        'Организационная структура и права доступа':
            ('Контур определяет, кто работает в системе и что ему разрешено. Модель доступа гибридная и трёхуровневая: '
             'базовая роль хранится в поле users_user.role, типовые права по должностям задаются в '
             'users_positionaccessrule, а точечные исключения для конкретного сотрудника — в users_useraccessoverride. '
             'Индивидуальные правила имеют наивысший приоритет. Подразделения образуют дерево через parent_id, '
             'а поле supervisor_id пользователя обеспечивает маршрутизацию согласования заявок «на непосредственного руководителя».',
             'org'),
        'Нормативно-справочная информация':
            ('Справочники — опора всех операционных контуров. Карточка актива references_asset универсальна для всех '
             'типов активов и содержит поля интеграции с 1С (source_1c_id, last_sync_at). Категории и группы товаров '
             'реализованы одной таблицей references_assetcategory с иерархией через parent_id. Единицы измерения и '
             'должности синхронизируются с зависимыми записями автоматически при переименовании.',
             'refs'),
        'Складской учёт':
            ('Контур разделяет текущее состояние (assets_warehousestock) и историю операций (assets_stockmovement). '
             'Журнал движений связывается с документом-основанием полиморфной ссылкой GenericForeignKey '
             '(document_type_id + document_id), что позволяет привязать движение к накладной, акту или заявке без '
             'отдельных полей под каждый тип. Правила алармов (stockalertrule) и их срабатывания (stockalertstate) '
             'обеспечивают контроль критических остатков с уведомлением ответственных.',
             'stock'),
        'Заявочный процесс':
            ('Заявка хранит бизнес-намерение сотрудника: что нужно, кому и почему. Позиции могут ссылаться как на '
             'конкретный актив, так и на группу товаров — это поддерживает пользовательский сценарий «интернет-магазина». '
             'Журнал согласования отделён от заявки: статус показывает текущее состояние, а requests_requestapproval '
             'хранит полную историю действий. Маршрут согласования настраивается по каждому виду заявки в requests_approvalstep.',
             'requests'),
        'Документооборот':
            ('Все документы построены по паттерну «заголовок + позиции» и наследуют общий набор полей абстрактного '
             'класса BaseDocument: номер (присваивается после финального подписания), дата, статус, автор и метки времени. '
             'Подписи вынесены в единую таблицу documents_documentsignature с полиморфной связью через ContentType — '
             'одна таблица обслуживает все пять типов документов. Члены комиссии также хранятся в одной таблице '
             'documents_commissionmember для актов, ходатайств и протоколов.',
             'documents'),
        'Уведомления и интеграции':
            ('Уведомления используют полиморфную связь related_content_type_id + related_object_id, что позволяет '
             'прикрепить уведомление к заявке, документу или складскому аларму без изменения структуры таблицы. '
             'Email-рассылка журналируется отдельно с фиксацией статуса и текста ошибки. Журнал integrations_synclog '
             'ведёт аудит каждого обмена с 1С: статус, количество созданных и обновлённых записей.',
             'notifications'),
    }

    num = 3
    for section_name, tables in ALL_SECTIONS:
        add_heading(doc, f'{num}. {section_name}', 1)
        intro, img_key = section_intro[section_name]
        add_para(doc, intro)
        add_image(doc, images[img_key], 16.5, f'Рисунок {num - 1}. Схема контура «{section_name}»')
        sub = 1
        for tname, purpose, fields in tables:
            add_heading(doc, f'{num}.{sub} Таблица {tname}', 3)
            add_para(doc, purpose, size=13)
            add_fields_table(doc, fields)
            doc.add_paragraph()
            sub += 1
        num += 1

    # ---------- Статусные модели ----------
    add_heading(doc, f'{num}. Статусные модели процессов', 1)
    add_para(doc, 'Заявки и документы являются процессными сущностями: их поле status определяет положение объекта '
                  'в жизненном цикле. Переходы выполняются сервисным слоем приложения и фиксируются в журналах '
                  'согласований и подписей, что обеспечивает полную воспроизводимость истории.')
    add_image(doc, images['statuses'], 16.5, f'Рисунок {num - 1}. Статусные модели заявок, документов и закреплений')
    num += 1

    # ---------- Ссылочная целостность ----------
    add_heading(doc, f'{num}. Стратегия ссылочной целостности', 1)
    add_para(doc, 'В схеме последовательно применяются три стратегии обработки удаления связанных записей:')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, name in enumerate(['Стратегия', 'Где применяется', 'Смысл']):
        run = hdr[i].paragraphs[0].add_run(name)
        set_font(run, 12, bold=True, color='FFFFFF')
        shade_cell(hdr[i], '1F3864')
    integrity = [
        ('PROTECT', 'Активы в позициях документов и заявок, категории активов, виды заявок, контрагенты и склады в накладных, единицы измерения в карточках.',
         'Справочник нельзя удалить, пока он используется в юридически значимых данных.'),
        ('CASCADE', 'Позиции документов и заявок, договоры контрагентов, индивидуальные права, срабатывания алармов.',
         'Дочерние записи не имеют смысла без родителя и удаляются вместе с ним.'),
        ('SET_NULL', 'Руководители, МОЛ, исполнители операций, склады в историческихзаписях, подразделения.',
         'Историческая запись сохраняется, даже если связанный объект удалён.'),
    ]
    for strat, where, sense in integrity:
        row = table.add_row().cells
        for i, val in enumerate([strat, where, sense]):
            run = row[i].paragraphs[0].add_run(val)
            set_font(run, 11)
    doc.add_paragraph()
    add_para(doc, 'Такое распределение защищает мастер-данные, позволяет операционным строкам следовать за заголовком '
                  'и гарантирует, что исторические журналы (движения, согласования, подписи) не разрушаются при '
                  'изменении справочников.')
    num += 1

    # ---------- Критичные бизнес-связи ----------
    add_heading(doc, f'{num}. Критичные бизнес-связи', 1)
    links = [
        ('Сотрудник → подразделение', 'users_user.department_id определяет принадлежность к подразделению; через неё строится видимость заявок для руководителя и коллег.'),
        ('Сотрудник → руководитель', 'users_user.supervisor_id хранит непосредственного руководителя и используется при маршрутизации этапа согласования с признаком requires_supervisor.'),
        ('Заявка → ответственные за выдачу', 'M2M-связь requests_assetrequest.issue_responsibles заполняется руководителем АХС при назначении исполнителей выдачи.'),
        ('Движение → документ-основание', 'assets_stockmovement.document_type_id + document_id полиморфно связывает фактическую операцию с накладной, актом или иным документом.'),
        ('Остаток → аларм', 'assets_stockalertstate соединяет правило контроля и конкретную строку остатка; уникальный индекс исключает дубли активных срабатываний.'),
        ('Документ → подписи', 'documents_documentsignature через ContentType обслуживает подписи всех пяти типов документов единообразно.'),
        ('Актив → 1С', 'references_asset.source_1c_id — уникальный внешний ключ интеграции; журнал загрузок ведётся в integrations_synclog.'),
    ]
    for title, text in links:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}. ')
        set_font(run, 14, bold=True)
        run2 = p.add_run(text)
        set_font(run2, 14)
        p.paragraph_format.space_after = Pt(6)
    num += 1

    # ---------- Ограничения и развитие ----------
    add_heading(doc, f'{num}. Известные ограничения и направления развития', 1)
    limits = [
        'Остатки: assets_warehousestock хранит одну строку остатка на актив (OneToOne). Для мультискладского учёта '
        '«один товар на нескольких складах» рекомендуется перейти на ForeignKey(asset) с уникальным индексом (asset_id, warehouse_id).',
        'Категории: references_assetcategory используется и как категория, и как группа. При росте иерархии целесообразно '
        'ввести явный признак уровня.',
        'Номера документов: поле number не уникально на уровне БД — номер генерируется внутри каждого типа документа. '
        'При требовании сквозной нумерации потребуется общий реестр номеров.',
        'Контрольные ограничения: для финансовых полей рекомендуется добавить ограничения БД quantity > 0, unit_price ≥ 0, total ≥ 0.',
        'Интеграция с 1С: ключом сопоставления должен служить source_1c_id (резервно — code); импорт должен быть идемпотентным, '
        'каждая загрузка — журналироваться в integrations_synclog с протоколом сверки.',
    ]
    for i, l in enumerate(limits, 1):
        add_para(doc, f'{i}. {l}', space_after=5)
    num += 1

    # ---------- Рекомендуемые индексы ----------
    add_heading(doc, f'{num}. Рекомендуемые индексы для производительности', 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, name in enumerate(['Таблица', 'Рекомендуемые индексы']):
        run = hdr[i].paragraphs[0].add_run(name)
        set_font(run, 12, bold=True, color='FFFFFF')
        shade_cell(hdr[i], '1F3864')
    idx = [
        ('requests_assetrequest', 'status; initiator_id; created_at; request_type_id'),
        ('requests_requestapproval', 'request_id; approver_id; created_at'),
        ('documents_* (заголовки)', 'status; created_by_id; created_at; date'),
        ('notifications_notification', 'составной: recipient_id + is_read + created_at'),
        ('assets_stockmovement', 'asset_id; warehouse_id; performed_at; movement_type'),
        ('references_asset', 'asset_type; category_id; group_id; source_1c_id'),
    ]
    for t, i_ in idx:
        row = table.add_row().cells
        run = row[0].paragraphs[0].add_run(t)
        set_font(run, 11)
        run = row[1].paragraphs[0].add_run(i_)
        set_font(run, 11)
    doc.add_paragraph()
    num += 1

    # ---------- Итог ----------
    add_heading(doc, f'{num}. Итоговая архитектурная оценка', 1)
    add_para(doc, 'Модель данных ИС «АСУ» последовательно разделяет справочники, процессы и фактические операции. '
                  'Сильные стороны схемы:')
    strengths = [
        'единая карточка актива для ТМЗ, ОС и НМА с полями интеграции с 1С;',
        'отдельный аудиторский журнал складских движений с привязкой к документу-основанию;',
        'отделённый от заявки журнал согласований с полной историей действий;',
        'универсальная полиморфная модель подписей документов;',
        'универсальная модель уведомлений через ContentType;',
        'гибкая трёхуровневая модель прав: роль, должностные правила, индивидуальные исключения.',
    ]
    for s in strengths:
        add_para(doc, f'— {s}', space_after=3)
    add_para(doc, 'Ключевая зона развития перед промышленной эксплуатацией — переход на мультискладской учёт остатков '
                  'с уникальностью «актив + склад», после чего система будет готова к массовой загрузке реальных '
                  'остатков из 1С в разрезе складов.')

    out = os.path.join(BASE_DIR, 'ASU_Описание_базы_данных.docx')
    doc.save(out)
    return out


def main():
    images = {
        'overview': diagram_overview(),
        'org': diagram_org(),
        'refs': diagram_refs(),
        'stock': diagram_stock(),
        'requests': diagram_requests(),
        'documents': diagram_documents(),
        'notifications': diagram_notifications(),
        'statuses': diagram_statuses(),
    }
    out = build_document(images)
    print('OK:', out)


if __name__ == '__main__':
    main()
